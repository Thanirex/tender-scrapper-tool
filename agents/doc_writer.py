import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_report(records: list, output_path: str) -> None:
    """
    Write a formatted Word report (.docx) from a list of tender records.
    Each record: {title, keyword, url, summary, files:[paths]}
    """
    doc = Document()
    _set_margins(doc)

    # Cover page
    h = doc.add_heading("Tender Intelligence Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    count_p = doc.add_paragraph(f"Total tenders: {len(records)}")
    count_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    for i, rec in enumerate(records, 1):
        _write_tender_section(doc, i, rec)
        if i < len(records):
            doc.add_page_break()

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def _write_tender_section(doc: Document, index: int, rec: dict) -> None:
    title = rec.get("title", "Untitled")
    keyword = rec.get("keyword", "")
    url = rec.get("url", "")
    summary = rec.get("summary", "No summary available.")
    files = rec.get("files", [])

    doc.add_heading(f"{index}. {title}", 2)

    # Source metadata line
    meta = doc.add_paragraph()
    meta.add_run("Keyword: ").bold = True
    meta.add_run(keyword + "    ")
    meta.add_run("Source: ").bold = True
    link_run = meta.add_run(url)
    link_run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph()

    # 5-field summary, parsed and rendered with bold labels
    _render_summary_fields(doc, summary)

    # List of downloaded attachments
    if files:
        doc.add_paragraph()
        doc.add_heading("Attached Documents", 4)
        for fpath in files:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(os.path.basename(fpath))


_MARKER_STYLES = {
    "VERIFIED":   {"text": "✓ VERIFIED",          "color": RGBColor(0x05, 0x96, 0x69)},  # green
    "EXTRACTED":  {"text": "~ EXTRACTED",          "color": RGBColor(0x0E, 0x7A, 0xBF)},  # blue
    "NOT_FOUND":  {"text": "✗ NOT FOUND",          "color": RGBColor(0xB9, 0x1C, 0x1C)},  # red
    "MISMATCH":   {"text": "⚠ MISMATCH: HIGH PRIORITY ERROR", "color": RGBColor(0xDC, 0x26, 0x26)},  # bright red
}

_MARKER_RE = re.compile(
    r"\[(VERIFIED|EXTRACTED|NOT_FOUND|MISMATCH(?::[^\]]*)?)\]\s*", re.IGNORECASE
)


def _render_summary_fields(doc: Document, summary: str) -> None:
    """Parse '1. Field Name: [MARKER] value' blocks and render with styled confidence badges."""
    field_pattern = re.compile(r"\d+\.\s*([^:\n]+):\s*(.*?)(?=\n\d+\.|$)", re.DOTALL)
    matches = field_pattern.findall(summary)

    if not matches:
        doc.add_paragraph(summary)
        return

    for label, raw_value in matches:
        label = label.strip()
        raw_value = raw_value.strip()

        # Detect marker
        m = _MARKER_RE.match(raw_value)
        marker_key = None
        value_text = raw_value

        if m:
            full_marker = m.group(1).upper()
            # Normalize MISMATCH variants
            if full_marker.startswith("MISMATCH"):
                marker_key = "MISMATCH"
            else:
                marker_key = full_marker
            value_text = raw_value[m.end():].strip()

        p = doc.add_paragraph()

        # Bold field label
        label_run = p.add_run(label + ":  ")
        label_run.bold = True

        # Confidence badge
        if marker_key and marker_key in _MARKER_STYLES:
            style = _MARKER_STYLES[marker_key]
            badge = p.add_run(f"[{style['text']}]  ")
            badge.bold = True
            badge.font.color.rgb = style["color"]
            badge.font.size = Pt(9)

        # Value text
        value_run = p.add_run(value_text)
        if marker_key == "MISMATCH":
            value_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            value_run.bold = True
        elif marker_key == "NOT_FOUND":
            value_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            value_run.italic = True

        # Extra spacing after mismatch to make it stand out
        if marker_key == "MISMATCH":
            doc.add_paragraph()


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
